"""CV download routes: /cv.pdf (the real formatted PDF) and /cv.docx (generated).

Also exposes helpers the chat uses to detect a "give me your CV" request and to
render the HTML download-buttons response.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from starlette.responses import FileResponse, Response, JSONResponse

ROOT = Path(__file__).resolve().parent
CV_PDF = ROOT / "docs" / "kaljuvee-julian-ai-engineer-2026.pdf"
CV_MD = ROOT / "prompts" / "shared" / "cv.md"

PDF_NAME = "Julian-Kaljuvee-CV.pdf"
DOCX_NAME = "Julian-Kaljuvee-CV.docx"


# ── Request detection + response HTML (used by chat) ────────────────────────

_CV_RE = re.compile(
    r"(?:\b(cv|resume|résumé|curriculum vitae)\b).*(?:\b(download|pdf|word|docx|copy|send|share|give|get|have|email)\b)"
    r"|(?:\b(download|send|share|give|get|copy of|can i (?:get|have)|email me)\b).*\b(cv|resume|résumé|curriculum vitae)\b"
    r"|^\s*(cv|resume|résumé)\s*\??\s*$",
    re.IGNORECASE,
)


def is_cv_request(message: str) -> bool:
    return bool(_CV_RE.search(message or ""))


def cv_response_html() -> str:
    """Markdown/HTML bubble with PDF + Word download buttons (rendered by marked.js)."""
    return (
        "Of course — here's Julian's CV. Grab it in whichever format you prefer:\n\n"
        '<div class="cv-downloads">'
        '<a class="cv-download-btn" href="/cv/pdf" target="_blank" rel="noopener">'
        '<span class="cv-dl-ic">PDF</span> Download PDF</a>'
        '<a class="cv-download-btn" href="/cv/docx">'
        '<span class="cv-dl-ic">DOC</span> Download Word</a>'
        "</div>\n\n"
        "You can also reach Julian directly at **kaljuvee@gmail.com** or on "
        "[LinkedIn](https://www.linkedin.com/in/juliankaljuvee/). "
        "Happy to answer anything else about his experience."
    )


# ── DOCX generation from cv.md ──────────────────────────────────────────────

def _build_docx() -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    md = CV_MD.read_text(encoding="utf-8") if CV_MD.exists() else "# Julian Kaljuvee"
    # cv.md's H1 carries an internal note — present a clean title in the download.
    md = md.replace("# Julian Kaljuvee — CV (source of truth)",
                    "# Julian Kaljuvee — Curriculum Vitae")
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    for raw in md.split("\n"):
        line = raw.rstrip()
        s = line.strip()
        if not s:
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("- ") or s.startswith("* "):
            _formatted(doc.add_paragraph(style="List Bullet"), s[2:])
        else:
            _formatted(doc.add_paragraph(), s)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _formatted(paragraph, text: str):
    """Render inline **bold** markdown into a docx paragraph."""
    for i, chunk in enumerate(re.split(r"(\*\*.+?\*\*)", text)):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        else:
            paragraph.add_run(chunk)


# ── Route registration ──────────────────────────────────────────────────────

def register_cv_routes(rt):
    @rt("/cv/pdf")
    def cv_pdf():
        if not CV_PDF.exists():
            return JSONResponse({"error": "CV not found"}, status_code=404)
        return FileResponse(str(CV_PDF), media_type="application/pdf", filename=PDF_NAME)

    @rt("/cv/docx")
    def cv_docx():
        data = _build_docx()
        return Response(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{DOCX_NAME}"'},
        )
